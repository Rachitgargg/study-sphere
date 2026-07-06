"""
Service for loading, parsing, and cleaning various file formats (PDF, DOCX, TXT).
Uses pypdf and python-docx for extraction, and cleans the text for downstream processes.
"""
import os
import re
from pypdf import PdfReader
import docx

class FileLoaderService:
    def __init__(self):
        pass
        
    def load_file(self, file_path: str) -> str:
        """
        Loads the contents of a file, extracts raw text, cleans it, and returns the string.
        Raises FileNotFoundError if file is missing, and ValueError if extraction fails or text is empty.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if os.path.getsize(file_path) == 0:
            raise ValueError("The uploaded file is empty.")
            
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        extracted_text = ""
        
        try:
            if ext == ".txt":
                extracted_text = self._load_txt(file_path)
            elif ext == ".pdf":
                extracted_text = self._load_pdf(file_path)
            elif ext == ".docx":
                extracted_text = self._load_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            # Raise ValueError to be caught by the FastAPI endpoint handler
            raise ValueError(f"Error parsing {ext.upper()} file: {str(e)}")
            
        cleaned_text = self.clean_text(extracted_text)
        if not cleaned_text:
            raise ValueError("No readable text could be extracted from the document.")
            
        return cleaned_text

    def _load_txt(self, file_path: str) -> str:
        """Loads plain text files, trying several encodings to avoid decode errors."""
        encodings = ["utf-8", "latin-1", "cp1252", "utf-16"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # Fallback if all else fails
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _load_pdf(self, file_path: str) -> str:
        """Loads PDF files using pypdf and joins page text."""
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)

    def _load_docx(self, file_path: str) -> str:
        """Loads DOCX files using python-docx, including paragraphs and tables."""
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts)

    def clean_text(self, text: str) -> str:
        """
        Normalizes and cleans the extracted text:
        - Removes null bytes
        - Standardizes line breaks to single newlines (\n)
        - Trims leading/trailing line whitespace
        - Collapses excessive blank lines (more than 2 consecutive newlines to 2)
        - Collapses multiple internal spaces/tabs to a single space
        """
        if not text:
            return ""
            
        # Remove null characters
        text = text.replace("\x00", "")
        
        # Standardize line endings
        text = re.sub(r"\r\n", "\n", text)
        text = re.sub(r"\r", "\n", text)
        
        # Process line by line to strip trailing whitespace and normalize spacing
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            # Strip trailing/leading spaces on each line
            line_stripped = line.strip()
            # Collapse multiple spaces or tabs inside a line to a single space
            line_cleaned = re.sub(r"[ \t]+", " ", line_stripped)
            cleaned_lines.append(line_cleaned)
            
        # Reassemble text
        text = "\n".join(cleaned_lines)
        
        # Normalize consecutive blank lines (limit to maximum 2 consecutive newlines)
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        return text.strip()
